from __future__ import annotations

"""
Lógica central para trabajar con secciones del dictamen:

- Leer rangos de celdas desde un libro de Excel.
- Convertirlos en párrafos + tablas.
- Generar documentos DOCX por sección o un documento final completo.

Este módulo NO tiene dependencias de interfaz gráfica ni de Streamlit;
puede ser usado tanto por la app web como por una futura app de escritorio
o por scripts de línea de comandos.
"""

import io
import json
from pathlib import Path
from typing import Any, Dict, Iterable, List, Tuple

import pandas as pd
from openpyxl import load_workbook
from openpyxl.workbook.workbook import Workbook
from docx import Document
from docx.document import Document as DocumentType
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm

from scripts.procesador_bloques import procesar_bloque_por_tipo


# ------------------------
# Carga de datos / config
# ------------------------

def cargar_rangos(config_path: Path) -> Dict[str, List[Dict[str, str]]]:
    """
    Carga el archivo JSON de rangos de hojas.

    Estructura esperada:
        {
          "Portada": "A1:Q60",
          "BG": "A1:Q120",
          ...
        }
    """
    if not config_path.exists():
        raise FileNotFoundError(f"No se encontró el archivo de rangos: {config_path}")

    with config_path.open("r", encoding="utf-8") as f:
        data = json.load(f)

    if not isinstance(data, dict):
        raise ValueError("El archivo de rangos debe contener un objeto JSON (dict) en la raíz.")

    rangos_normalizados: Dict[str, List[Dict[str, str]]] = {}

    for hoja, valor in data.items():
        # Formato legacy: una sola cadena con uno o varios subrangos
        # por ejemplo: "A1:F40;A41:F60"
        if isinstance(valor, str):
            rangos_normalizados[hoja] = [
                {
                    "rango": valor,
                    "tipo": "texto_normal",
                }
            ]
            continue

        # Nuevo formato: lista de bloques {rango, tipo}
        if isinstance(valor, list):
            bloques: List[Dict[str, str]] = []
            for bloque in valor:
                if isinstance(bloque, dict):
                    rango = str(bloque.get("rango", "")).strip()
                    if not rango:
                        continue
                    tipo = str(bloque.get("tipo", "texto_normal")).strip() or "texto_normal"
                    bloques.append({"rango": rango, "tipo": tipo})
                elif isinstance(bloque, str):
                    # Soporte de compatibilidad: lista de strings
                    bloques.append({"rango": bloque.strip(), "tipo": "texto_normal"})
            if bloques:
                rangos_normalizados[hoja] = bloques
            continue

        raise ValueError(
            f"Formato no soportado en rangos_hojas para la hoja '{hoja}': "
            f"se esperaba cadena o lista, se recibió {type(valor).__name__}"
        )

    return rangos_normalizados


def cargar_workbook(excel_path: Path) -> Workbook:
    """
    Carga el libro de Excel en modo solo datos.
    """
    if not excel_path.exists():
        raise FileNotFoundError(f"No se encontró el archivo de Excel: {excel_path}")

    return load_workbook(excel_path, data_only=True)


def cargar_formatos(formatos_path: Path) -> Dict[str, Any]:
    """
    Carga el archivo JSON de formatos de hojas.

    Estructura esperada:
        {
          "mapa_hojas": { "BG": "estado_financiero", ... },
          "perfiles": { "estado_financiero": { ... }, ... }
        }
    """
    if not formatos_path.exists():
        raise FileNotFoundError(f"No se encontró el archivo de formatos: {formatos_path}")

    with formatos_path.open("r", encoding="utf-8") as f:
        data = json.load(f)

    if not isinstance(data, dict):
        raise ValueError("El archivo de formatos debe contener un objeto JSON (dict).")

    return data  # type: ignore[return-value]



# ------------------------
# Extracción de secciones
# ------------------------

def extraer_seccion_desde_hoja(
    wb: Workbook,
    sheet_name: str,
    rango: str,
) -> Tuple[List[str], List[pd.DataFrame]]:
    """
    A partir de una hoja y un rango, devuelve:
      - lista de párrafos (strings)
      - lista de tablas (cada tabla es un DataFrame)

    La lógica de separación es:
      - Filas completamente vacías => separador de párrafos / "salto"
      - Filas con números => consideradas parte de una tabla
      - Filas sin números => consideradas texto corrido
    """
    # Soporta múltiples rangos separados por ';', por ejemplo:
    # "A1:Q50;A60:Q120"
    if ";" in rango:
        parrafos_totales: List[str] = []
        tablas_totales: List[pd.DataFrame] = []

        for sub_rango in [r.strip() for r in rango.split(";") if r.strip()]:
            p, t = extraer_seccion_desde_hoja(wb, sheet_name, sub_rango)
            # Separador suave entre bloques de rango
            if parrafos_totales:
                parrafos_totales.append("")
            parrafos_totales.extend(p)
            tablas_totales.extend(t)

        return parrafos_totales, tablas_totales

    if sheet_name not in wb.sheetnames:
        raise KeyError(f"La hoja '{sheet_name}' no existe en el libro de Excel.")

    hoja = wb[sheet_name]
    celdas = hoja[rango]

    parrafos: List[str] = []
    tablas: List[pd.DataFrame] = []

    tabla_actual: List[List[object]] = []

    for row in celdas:
        valores = [c.value for c in row]

        # Fila completamente vacía => separador / posible fin de tabla
        if all(v is None for v in valores):
            if tabla_actual:
                df = pd.DataFrame(tabla_actual)
                tablas.append(df)
                tabla_actual = []
            parrafos.append("")  # separador
            continue

        # Heurística: ¿parece fila de tabla?
        fila_es_tabla = any(isinstance(v, (int, float)) for v in valores if v is not None)

        if fila_es_tabla:
            tabla_actual.append(valores)
        else:
            # Si veníamos acumulando tabla, la cerramos
            if tabla_actual:
                df = pd.DataFrame(tabla_actual)
                tablas.append(df)
                tabla_actual = []

            texto = " ".join([str(v) for v in valores if v not in (None, "")])
            if texto.strip():
                parrafos.append(texto)

    # Tabla pendiente al final
    if tabla_actual:
        df = pd.DataFrame(tabla_actual)
        tablas.append(df)

    return parrafos, tablas


def _create_doc_from_template(plantilla_path: Path) -> DocumentType:
    """
    Crea un Document basado en la plantilla, pero
    limpiando todo el contenido del cuerpo (párrafos y tablas).

    De esta forma reutilizamos estilos, encabezados/pies y configuración,
    sin arrastrar el contenido fijo de la plantilla en cada sección.
    """
    doc: DocumentType = Document(plantilla_path)

    # Eliminar párrafos existentes
    for p in list(doc.paragraphs):
        p_element = p._element
        parent = p_element.getparent()
        if parent is not None:
            parent.remove(p_element)

    # Eliminar tablas existentes
    for t in list(doc.tables):
        t_element = t._element
        parent = t_element.getparent()
        if parent is not None:
            parent.remove(t_element)

    return doc


def _cache_model_tables(doc: DocumentType) -> Dict[str, Any]:
    """
    Finds all tables in the document that look like model tables
    (i.e., have an ID like [[...]] in their first cell) and returns
    a dictionary mapping the ID to the table's XML element and column widths.
    """
    cache = {}
    for table in doc.tables:
        if table.rows and table.columns and table.cell(0, 0).text.strip().startswith("[["):
            model_id = table.cell(0, 0).text.strip()
            if model_id.endswith("]]"):
                # Cache both the XML element and the measured column widths
                cache[model_id] = {
                    "xml": table._tbl,
                    "widths": [col.width for col in table.columns],
                }
    return cache


def generar_docx_seccion_a_archivo(
    wb: Workbook,
    sheet_name: str,
    bloques: List[Dict[str, str]],
    plantilla_path: Path,
    destino: Path,
    formatos: Dict[str, Any] | None = None,
) -> None:
    """
    Genera un DOCX de una sola hoja usando la arquitectura basada en bloques.
    """
    if not plantilla_path.exists():
        raise FileNotFoundError(f"No se encontró la plantilla de Word: {plantilla_path}")

    template_doc = Document(plantilla_path)
    model_tables = _cache_model_tables(template_doc)

    # Create a new doc for the section, but based on the original template
    doc = _create_doc_from_template(plantilla_path)

    for bloque in bloques:
        procesar_bloque_por_tipo(wb, sheet_name, bloque, doc, formatos, model_tables)

    destino.parent.mkdir(parents=True, exist_ok=True)
    with destino.open("wb") as f:
        doc.save(f)


def generar_docx_final_en_memoria(
    wb: Workbook,
    rangos: Dict[str, List[Dict[str, str]]],
    plantilla_path: Path,
    orden: Iterable[str] | None = None,
    formatos: Dict[str, Any] | None = None,
) -> io.BytesIO:
    """
    Genera un DOCX final combinando múltiples secciones en memoria.

    - `rangos`: mapping hoja -> lista de bloques [{rango, tipo}]
    - `orden`: orden explícito de hojas; si es None, se usa el orden de `rangos`.
    """
    from docxcompose.composer import Composer

    if not plantilla_path.exists():
        raise FileNotFoundError(f"No se encontró la plantilla de Word: {plantilla_path}")

    template_doc = Document(plantilla_path)
    model_tables = _cache_model_tables(template_doc)

    # Create a clean base document for the composer
    base = _create_doc_from_template(plantilla_path)
    composer = Composer(base)

    # Determinar orden efectivo
    if orden is None:
        orden_efectivo = [h for h in rangos.keys() if h in wb.sheetnames]
    else:
        orden_efectivo = [h for h in orden if h in rangos and h in wb.sheetnames]

    for i, sheet_name in enumerate(orden_efectivo):
        bloques = rangos.get(sheet_name, [])
        if not bloques:
            continue

        doc_sec = _create_doc_from_template(plantilla_path)
        for bloque in bloques:
            procesar_bloque_por_tipo(wb, sheet_name, bloque, doc_sec, formatos, model_tables)

        composer.append(doc_sec)

        # Añadir un salto de página después de cada sección, con la lógica específica.
        # Solo saltar después de las hojas con índice 1 a 7 (la 2da a la 8va hoja)
        # y no si es la última hoja del documento.
        if i >= 1 and i <= 7 and i < len(orden_efectivo) - 1:
            composer.doc.add_page_break()

    out_buf = io.BytesIO()
    composer.save(out_buf)
    out_buf.seek(0)
    return out_buf


def generar_docx_final_a_archivo(
    wb: Workbook,
    rangos: Dict[str, List[Dict[str, str]]],
    plantilla_path: Path,
    destino: Path,
    orden: Iterable[str] | None = None,
    formatos: Dict[str, Any] | None = None,
) -> None:
    """
    Genera y guarda el DOCX final en disco.
    """
    buffer = generar_docx_final_en_memoria(wb, rangos, plantilla_path, orden=orden, formatos=formatos)
    destino.parent.mkdir(parents=True, exist_ok=True)
    with destino.open("wb") as f:
        f.write(buffer.read())
