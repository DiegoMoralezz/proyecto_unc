from __future__ import annotations

import copy
from pathlib import Path
from typing import Any, Dict, Iterable, List

import pandas as pd
from openpyxl.workbook.workbook import Workbook
from docx import Document
from docx.document import Document as DocumentType
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt


def _get_paragraph_alignment(name: str | None) -> int | None:
    if not name:
        return None
    name = name.lower()
    if name == "left":
        return WD_ALIGN_PARAGRAPH.LEFT
    if name == "right":
        return WD_ALIGN_PARAGRAPH.RIGHT
    if name == "center":
        return WD_ALIGN_PARAGRAPH.CENTER
    if name == "justify":
        return WD_ALIGN_PARAGRAPH.JUSTIFY
    return None


def _leer_rango_celdas(wb: Workbook, sheet_name: str, rango: str):
    if sheet_name not in wb.sheetnames:
        raise KeyError(f"La hoja '{sheet_name}' no existe en el libro de Excel.")
    hoja = wb[sheet_name]
    return hoja[rango]


def procesar_bloque_por_tipo(
    wb: Workbook,
    sheet_name: str,
    bloque: Dict[str, Any], # Changed to Any to support 'contenido'
    doc: DocumentType,
    formatos: Dict[str, Any] | None,
    model_tables_cache: Dict[str, Any] | None = None,
) -> None:
    """
    Aplica el procesamiento adecuado según bloque["tipo"], usando
    la configuración declarativa en formatos["tipos"].
    Maneja bloques con 'rango' (legacy) y con 'contenido' (automático).
    """
    tipo = bloque.get("tipo", "").strip()
    if not tipo:
        return

    # Extraer configuración de formato para este tipo de bloque
    config_tipo: Dict[str, Any] = {}
    if formatos is not None:
        tipos_cfg = formatos.get("tipos", {}) or {}
        config_tipo = tipos_cfg.get(tipo, {}) or {}

    # --- Lógica de bifurcación: por contenido directo o por rango ---
    contenido_directo = bloque.get("contenido")
    rango = bloque.get("rango", "").strip()

    if contenido_directo is not None:
        # --- NUEVO: Procesar bloque con contenido directo ---
        if tipo.startswith("tabla_"):
            if isinstance(contenido_directo, pd.DataFrame):
                df_rows = contenido_directo.values.tolist()
                _procesar_tabla_directo(doc, df_rows, config_tipo, model_tables_cache)
            else:
                print(f"WARN: Bloque tipo tabla '{tipo}' no contiene un DataFrame. Se ignora.")
        else: # Tratar como cualquier tipo de texto
            texto = str(contenido_directo)
            _procesar_texto_directo(doc, texto, config_tipo)
    elif rango:
        # --- LEGACY: Procesar bloque con rango ---
        if tipo.startswith("tabla_"):
            excel_rows = [[c.value for c in row] for row in _leer_rango_celdas(wb, sheet_name, rango)]
            _procesar_tabla_directo(doc, excel_rows, config_tipo, model_tables_cache)
        elif tipo.startswith("texto_sangria"):
            _procesar_texto_rango(wb, sheet_name, rango, doc, config_tipo)
        elif tipo.startswith("viñetas") or tipo.startswith("vinyetas"):
            _procesar_viñetas_rango(wb, sheet_name, rango, doc, config_tipo)
        elif tipo.startswith("titulo_"):
            _procesar_titulo_rango(wb, sheet_name, rango, doc, config_tipo)
        elif tipo == "num_notas":
            _procesar_numero_nota_rango(wb, sheet_name, rango, doc, config_tipo)
        else: # Fallback: texto normal
            _procesar_texto_rango(wb, sheet_name, rango, doc, config_tipo)
    else:
        # El bloque no tiene ni 'contenido' ni 'rango', no se puede procesar.
        return


def _aplicar_parrafo_config(p, config: Dict[str, Any]) -> None:
    style_name = config.get("style")
    if style_name:
        candidates = [s.strip() for s in str(style_name).split(",") if s.strip()]
        for candidate in candidates:
            try:
                p.style = candidate
                break
            except Exception:
                continue

    align_name = config.get("align")
    align_val = _get_paragraph_alignment(align_name)
    if align_val is not None:
        p.alignment = align_val

    first_line_indent = config.get("first_line_indent")
    if isinstance(first_line_indent, (int, float)):
        try:
            p.paragraph_format.first_line_indent = Cm(float(first_line_indent))
        except Exception:
            pass

# --- NUEVAS funciones para procesar contenido directo ---

def _procesar_texto_directo(doc: DocumentType, texto: str, config_tipo: Dict[str, Any]) -> None:
    """Procesa un bloque de texto simple a partir de un string de contenido."""
    if not texto.strip():
        # Si el contenido es solo espacios en blanco, podría ser intencional
        p = doc.add_paragraph()
        _aplicar_parrafo_config(p, config_tipo)
        return

    # Tratar saltos de línea en el contenido como párrafos separados
    for line in texto.split('\n'):
        p = doc.add_paragraph(line)
        _aplicar_parrafo_config(p, config_tipo)


def _procesar_tabla_directo(
    doc: DocumentType,
    filas_datos: List[List[Any]],
    config_tipo: Dict[str, Any],
    model_tables_cache: Dict[str, Any] | None = None,
) -> None:
    """Procesa un bloque de tabla a partir de una lista de listas de datos."""
    table_model_id = config_tipo.get("table_model_id")
    if table_model_id and model_tables_cache and table_model_id in model_tables_cache:
        model_data = model_tables_cache[table_model_id]
        _crear_tabla_clonada(doc, model_data["xml"], model_data["widths"], filas_datos, config_tipo)
    else:
        # Fallback a un método simple si no hay modelo de tabla
        _crear_tabla_desde_datos(doc, filas_datos, config_tipo)


# --- Funciones LEGACY adaptadas (renombradas a _*_rango) ---

def _procesar_texto_rango(
    wb: Workbook, sheet_name: str, rango: str, doc: DocumentType, config_tipo: Dict[str, Any]
) -> None:
    celdas = _leer_rango_celdas(wb, sheet_name, rango)
    for row in celdas:
        valores = [c.value for c in row]
        if all(v is None for v in valores):
            p = doc.add_paragraph()
            _aplicar_parrafo_config(p, config_tipo)
            continue
        texto = " ".join(str(v) for v in valores if v not in (None, ""))
        if not texto.strip():
            continue
        p = doc.add_paragraph(texto)
        _aplicar_parrafo_config(p, config_tipo)


def _procesar_viñetas_rango(
    wb: Workbook, sheet_name: str, rango: str, doc: DocumentType, config_tipo: Dict[str, Any]
) -> None:
    celdas = _leer_rango_celdas(wb, sheet_name, rango)
    for row in celdas:
        valores = [c.value for c in row]
        texto = " ".join(str(v) for v in valores if v not in (None, ""))
        if not texto.strip(): continue
        p = doc.add_paragraph(texto)
        _aplicar_parrafo_config(p, config_tipo)


def _procesar_titulo_rango(
    wb: Workbook, sheet_name: str, rango: str, doc: DocumentType, config_tipo: Dict[str, Any]
) -> None:
    celdas = _leer_rango_celdas(wb, sheet_name, rango)
    textos = [str(c.value) for row in celdas for c in row if c.value is not None]
    texto = " ".join(textos).strip()
    if not texto: return

    if config_tipo.get("page_break_before"):
        doc.add_page_break()
    
    _procesar_texto_directo(doc, texto, config_tipo)


def _procesar_numero_nota_rango(
    wb: Workbook, sheet_name: str, rango: str, doc: DocumentType, config_tipo: Dict[str, Any]
) -> None:
    celdas = _leer_rango_celdas(wb, sheet_name, rango)
    first_cell_value = celdas[0][0].value
    if first_cell_value is None: return
    
    texto = str(first_cell_value).strip()
    if not texto: return

    style_name = config_tipo.get("style")
    p = doc.add_paragraph(texto, style=style_name if style_name and style_name in doc.styles else None)
    _aplicar_parrafo_config(p, config_tipo)


def _crear_tabla_clonada(
    doc: DocumentType,
    model_table_xml: Any,
    column_widths: List[int],
    excel_rows: List[List[Any]],
    config_tipo: Dict[str, Any],
) -> None:
    """
    Clona una tabla modelo, la rellena con datos y la inserta en la posición
    actual del documento para mantener el orden de los bloques.
    """
    # Importación local para evitar dependencia a nivel de módulo si no se usa
    from docx.table import _Cell

    # 1. Añadir un párrafo marcador temporal para saber dónde insertar la tabla.
    marker_paragraph = doc.add_paragraph()
    
    # 2. Clonar el XML de la tabla modelo
    new_tbl_xml = copy.deepcopy(model_table_xml)

    # 3. Insertar el XML de la tabla justo después del párrafo marcador.
    marker_paragraph._p.addnext(new_tbl_xml)
    
    # 4. Obtener una referencia al objeto de tabla recién insertado.
    new_table = doc.tables[-1]

    # 5. Eliminar el párrafo marcador que ya no es necesario.
    p_element = marker_paragraph._p
    p_element.getparent().remove(p_element)

    # 6. Eliminar la fila del modelo si todavía existe
    if new_table.rows and new_table.cell(0, 0).text.strip().startswith("[["):
        new_table._tbl.remove(new_table.rows[0]._tr)

    # 7. Forzar el ancho de las columnas para asegurar la consistencia
    if column_widths and len(new_table.columns) == len(column_widths):
        for i, width in enumerate(column_widths):
            new_table.columns[i].width = width

    # --- Lógica de población ---
    header_rows_count = max(int(config_tipo.get("header_rows", 1)), 0)

    # Separar encabezados y datos del Excel
    excel_header_rows = excel_rows[:header_rows_count]
    excel_rows_data = excel_rows[header_rows_count:]

    # Sobrescribir los encabezados en la tabla de Word de forma segura
    for i, excel_header_row in enumerate(excel_header_rows):
        if i < len(new_table.rows) and i < header_rows_count:
            table_header_row = new_table.rows[i]
            # Acceso de bajo nivel para evitar el error 'no tr above'
            tr = table_header_row._tr
            for j, cell_data in enumerate(excel_header_row):
                if j < len(tr.tc_lst):
                    tc = tr.tc_lst[j]
                    cell = _Cell(tc, table_header_row)
                    
                    # Aplicar formato inteligente para años en encabezados
                    if isinstance(cell_data, float) and cell_data.is_integer():
                        text_to_write = str(int(cell_data))
                    else:
                        text_to_write = str(cell_data) if cell_data is not None else ""
                    
                    if not cell.paragraphs:
                        p = cell.add_paragraph()
                    else:
                        p = cell.paragraphs[0]
                    
                    if not p.runs:
                        p.add_run(text_to_write)
                    else:
                        p.runs[0].text = text_to_write
                        # Limpiar runs extra
                        for k in range(len(p.runs) - 1, 0, -1):
                            p._p.remove(p.runs[k]._r)

    # Ajustar filas de datos
    if config_tipo.get("trim_leading_empty_rows"):
        while excel_rows_data and _fila_vacia(excel_rows_data[0]):
            excel_rows_data = excel_rows_data[1:]
    
    if len(new_table.rows) <= header_rows_count: return

    model_data_rows_count = len(new_table.rows) - header_rows_count
    excel_data_rows_count = len(excel_rows_data)

    if excel_data_rows_count > model_data_rows_count:
        rows_to_add = excel_data_rows_count - model_data_rows_count
        template_row_xml = copy.deepcopy(new_table.rows[-1]._tr)
        for _ in range(rows_to_add):
            new_table._tbl.append(copy.deepcopy(template_row_xml))
    elif excel_data_rows_count < model_data_rows_count:
        rows_to_remove = model_data_rows_count - excel_data_rows_count
        for _ in range(rows_to_remove):
            new_table._tbl.remove(new_table.rows[-1]._tr)

    # Poblar la tabla con los datos
    for i, excel_row_data in enumerate(excel_rows_data):
        table_row_index = i + header_rows_count
        if table_row_index < len(new_table.rows):
            table_row = new_table.rows[table_row_index]
            for j, cell_data in enumerate(excel_row_data):
                if j < len(table_row.cells):
                    cell = table_row.cells[j]
                    texto = _formatear_celda_tabla(cell_data, table_row_index, j, config_tipo)
                    
                    # 0. Asegurar que tenemos el párrafo de la celda actual
                    if not cell.paragraphs:
                        p = cell.add_paragraph()
                    else:
                        p = cell.paragraphs[0]

                    # 1. Detectar fuente y tamaño de la plantilla (herencia de formato)
                    template_font = None
                    template_size = None
                    if p.runs:
                        template_font = p.runs[0].font.name
                        if p.runs[0].font.size:
                            template_size = p.runs[0].font.size.pt
                    
                    # 2. UNIFICAR Y ESCRIBIR EL CONTENIDO DEL EXCEL
                    if not p.runs:
                        new_run = p.add_run(texto)
                    else:
                        new_run = p.runs[0]
                        new_run.text = texto
                        # Limpiar cualquier otro run sobrante
                        for k in range(len(p.runs) - 1, 0, -1):
                            p._p.remove(p.runs[k]._r)
                    
                    # 3. Forzar fuente (JSON > Plantilla > Trebuchet MS)
                    font_name_cfg = config_tipo.get("font_name")
                    f_name = font_name_cfg or template_font or "Trebuchet MS"

                    if f_name:
                        new_run.font.name = f_name
                        from docx.oxml.ns import qn
                        rPr = new_run._element.get_or_add_rPr()
                        rFonts = rPr.get_or_add_rFonts()
                        rFonts.set(qn('w:ascii'), f_name)
                        rFonts.set(qn('w:hAnsi'), f_name)
                        rFonts.set(qn('w:cs'), f_name)

                    # 4. Forzar tamaño (JSON > Plantilla > 10.0)
                    column_font_sizes = config_tipo.get("column_font_size", {})
                    col_size = None
                    if isinstance(column_font_sizes, list) and j < len(column_font_sizes):
                        col_size = column_font_sizes[j]
                    elif isinstance(column_font_sizes, dict):
                        col_size = column_font_sizes.get(j) or column_font_sizes.get(j + 1)
                    
                    f_size = col_size or config_tipo.get("font_size") or template_size or 10.0
                    
                    try:
                        new_run.font.size = Pt(float(f_size))
                    except: pass


def _formatear_celda_tabla(
    cell_data: Any,
    row_index: int,
    col_index: int,
    config_tipo: Dict[str, Any],
) -> str:
    """
    Formatea el valor de una celda para su inserción en la tabla Word,
    manejando NaNs y aplicando formatos numéricos contextuales.
    """
    # 1. Manejar valores nulos o NaN de pandas
    if cell_data is None or pd.isna(cell_data):
        return ""

    # 2. Manejar valores numéricos (int y float)
    if isinstance(cell_data, (int, float)):
        # Si es cero, devolver un guion
        if cell_data == 0:
            return "-"

        # Revisar si hay una regla de formato específica en la configuración
        number_format_type = _resolver_formato_numerico(config_tipo, col_index)

        # --- NUEVA LÓGICA DE PORCENTAJE ---
        if number_format_type == "percentage":
            return f"{cell_data:.2%}"

        # Lógica especial para las primeras dos filas (si es necesario)
        if row_index < 2:
            val = int(cell_data)
            return f"({abs(val)})" if val < 0 else str(val)

        # Lógica estándar para el resto de filas
        is_integer_float = isinstance(cell_data, float) and cell_data.is_integer()
        
        if isinstance(cell_data, int) or is_integer_float:
            # Formatear como entero con separador de miles
            val = int(cell_data)
            return f"({abs(val):,})" if val < 0 else f"{val:,}"
        else:
            # Formatear como flotante con 2 decimales y separador de miles
            return f"({abs(cell_data):,.2f})" if cell_data < 0 else f"{cell_data:,.2f}"

    # 3. Fallback para cualquier otro tipo de dato (como strings)
    return str(cell_data)


def _resolver_formato_numerico(config_tipo: Dict[str, Any], col_index: int) -> Any:
    """
    Devuelve el formato numérico efectivo para una columna dada:
    - column_number_format puede ser lista o dict (0 o 1-based)
    - fallback a number_format global
    """
    col_cfg = config_tipo.get("column_number_format")
    if isinstance(col_cfg, list):
        if col_index < len(col_cfg):
            fmt = col_cfg[col_index]
            if fmt is not None:
                return fmt
    elif isinstance(col_cfg, dict):
        fmt = col_cfg.get(col_index)
        if fmt is None:
            fmt = col_cfg.get(col_index + 1)
        if fmt is not None:
            return fmt

    return config_tipo.get("number_format")


def _fila_vacia(row: List[Any]) -> bool:
    """Devuelve True si toda la fila es None o cadena vacia."""
    return all(cell is None or (isinstance(cell, str) and not cell.strip()) for cell in row)


def _crear_tabla_desde_datos(
    doc: DocumentType,
    filas_datos: List[List[Any]],
    config_tipo: Dict[str, Any],
) -> None:
    """Crea una tabla simple en Word a partir de una lista de listas (sin leer de Excel)."""
    if not filas_datos: return

    num_cols = len(filas_datos[0])
    table = doc.add_table(rows=0, cols=num_cols)
    table.style = config_tipo.get("table_style", "Table Grid")

    for fila in filas_datos:
        row_cells = table.add_row().cells
        for j, cell_val in enumerate(fila):
            row_cells[j].text = str(cell_val) if cell_val is not None else ""
